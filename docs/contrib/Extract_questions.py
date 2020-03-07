"""Contributed by Tianxin Shao

Sometimes we need Copy Questions from Assignx.docx to R_template.rmd. I
wrote one short Python application to easy this tedious things. This
application need "python_docx" module, help it useful, Thanks!
"""

pip install python_docx


import docx

def getQuestions(filename):
    doc = docx.Document(filename)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    print("# ",para.text)

def main():
    getQuestions("Assign2.docx")

if __name__ == "__main__":
    main()
