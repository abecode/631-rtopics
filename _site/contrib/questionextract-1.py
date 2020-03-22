"""Contributed by Anna Reasner

Builds on previous script, Extract_questions.py by adding a few
features to the script to make it automatically write this to rmd.

Note: (For assignment 3 you have to remove the ◦ bullet points in the
doc for the script to work)

"""

 

import docx

def getQuestions(filename):
    doc = docx.Document(filename)
    output = open("hw#-AnnaReasner.Rmd","w") 
    texty = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    ## print("# ",para.text)
                    texty = "# " + para.text
                    output.write(str(texty))
                    output.write("\n\n\n")
    

def main():
    getQuestions("#####.docx")

main()
